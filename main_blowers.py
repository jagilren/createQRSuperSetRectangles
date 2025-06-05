#main_blowers.py genera en Word los QR tipo Adhesive, está en orientacion landscape
import os
import qrcode
from PIL import Image, ImageDraw, ImageFont
import time
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def create_WordDocument():
    # Define the folder containing images
    image_folder = "URLS"

    # Get all image file paths
    image_files = [os.path.join(image_folder, f) for f in os.listdir(image_folder) if
                   f.lower().endswith(('png', 'jpg', 'jpeg'))]

    # Create a new Word document
    doc = Document()

    # Set document to landscape mode
    section = doc.sections[0]
    #Hoja en landscape direction Activar para Blowers y CCM
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width  # Swap width & height

    # Create a table with 2 columns
    table = doc.add_table(rows=0, cols=2)

    # Iterate over images in pairs
    for i in range(0, len(image_files), 2):
        row = table.add_row().cells  # Add a new row

        # Get the paragraph inside the cell
        paragraph = row[0].paragraphs[0]
        # Center-align the paragraph
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add first image
        #Para Equipos e Instrumentos
        #row[0].paragraphs[0].add_run().add_picture(image_files[i], width=Inches(3.3),height=Inches(2.16))
        #Para Blowers y CCM
        row[0].paragraphs[0].add_run().add_picture(image_files[i], width=Inches(3.7), height=Inches(2.2))


        # Add second image if available
        paragraph = row[1].paragraphs[0]
        # Center-align the paragraph
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if i + 1 < len(image_files):
            #Para Equipos e Instrumentos
            #row[1].paragraphs[0].add_run().add_picture(image_files[i + 1], width=Inches(3.3), height=Inches(2.16))
            #Para Blowers y CCM
            row[1].paragraphs[0].add_run().add_picture(image_files[i + 1], width=Inches(3.7), height=Inches(2.2))

        # Set text in blank row (needed to control font size)

        # Add an extra blank row after each image row
        blank_row = table.add_row().cells

        for cell in blank_row:
            # Remove space after the paragraph
            p_format = paragraph.paragraph_format
            p_format.space_after = Pt(0)  # Set space after to 0 pt

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(" ")  # Add a space so font can be applied
            #Para Equipos e Instrumentos
            #run.font.size = Pt(3)  # Set font size to 3 pt
            #Para Blowers y CCM
            run.font.size = Pt(7)  # Set font size to 7 pt


            #blank_row[0].text = ""  # Left column blank
            #blank_row[1].text = ""  # Right column blank

    # Save the Word document
    doc.save("Images_Table.docx")
    print("Word document created successfully!")
    return True

def create_tag_text_logoRPCI(qr_image_with_TAG_Logo,qr_height,qr_width):
    # Open the image file
    image_logo = Image.open("LOGO_RPCI.jpg")

    # Resize the image to 200x200 pixels
    resized_image_logo = image_logo.resize((int(white_rect_width/(2.5)), int(white_rect_width/3.0))) #Para respetar la relación de aspecto del logo de W/H = 0.91
    #Pone logo grande de RPCI a la derecha del QR y encima del Texto del TAG
    logo_x = white_rect_width/2 + 15
    logo_y = top_margin + 35 #Para que el logo de RPCI quede centrado con respecto al QR

    # Paste QR  Code created above  onto background
    qr_image_with_TAG_Logo.paste(resized_image_logo, (int(logo_x), int(logo_y)))
    return qr_image_with_TAG_Logo


def create_TagTex_at_Bottom(qr_image):
    # Function for add text to bottom of QR Code
    # Create image with white background

    image = Image.new("RGB", (white_rect_width, white_rect_height), BACKGROUND_COLOR)
    draw = ImageDraw.Draw(image)

    # Calculate QR position
    qr_x = int(white_rect_width *0.0000)  #(white_rect_width - qr_size) // 2
    qr_y = top_margin

    # Paste QR  Code created above  onto background
    image.paste(qr_image, (qr_x, qr_y))

    # Calculate text position and size
    try:
        font_path = "C:/Windows/Fonts/arialbd.ttf"  # Adjust if using Mac/Linux
        # Para Equipos e instrumentos
        #font_size = 0.5
        #Para Blower y CCM
        font_size = 1


        # Find optimal font size to match QR width
        while True:
            font = ImageFont.truetype(font_path, font_size)
            text_width = font.getlength(TAG)
            #Ancho También debe ser proporcional al # de carácteres.
            #Debemos establecer otro condicional
            if text_width > BASE_WIDTH/1.64 * 1.1 and len(TAG)>12:
                break

            elif text_width > BASE_WIDTH/1.64 * 0.85 and len(TAG)>6:
                break
            elif text_width > BASE_WIDTH / 1.64 * 0.4 and len(TAG) <= 6:
                break
            font_size += 1

        font_size = max(font_size - 1, 1)
        font = ImageFont.truetype(font_path, font_size)

        # Step1  Define Initial height of Text (TAG)
        text_bbox = font.getbbox(TAG)
        text_width = text_bbox[2] -text_bbox[0]
        if  len(TAG)>12:
            text_height = int((text_bbox[3] - text_bbox[1])*2) #Se multiplica por un factor reductor para que el texto se un poco más alto
        else:
            text_height = int((text_bbox[3] - text_bbox[
                1])*1)  # Se multiplica por un factor reductor para que el texto no sea tan alto

        # Create a transparent image sized to the text's bounding box
        text_image = Image.new("RGBA", (text_width, text_height), (0, 0, 0, 0))
        text_draw = ImageDraw.Draw(text_image)

        # Draw the text onto the transparent image (offset by bbox's left/top)
        text_draw.text((-text_bbox[0], -text_bbox[1]), TAG, font=font, fill="black")


        # STEP 2: Resize the text to reduce height (keep original width), ignore proportional aspect ratio of font
        if len(TAG)<12: #Para Textos que no son muy largos es necesario reducir la Heigth
            scale_factor = 0.65
        else: #Para Textos largos no es necesario reducir la Heigth
            scale_factor = 1

        new_height = int(text_height * scale_factor)
        resized_text = text_image.resize(
            (text_width, new_height),  # Keep width, reduce height
            resample=Image.Resampling.LANCZOS # High-quality resampling
        )

        # STEP 3: Paste the resized text onto the base image
        # --------------------------------------------------

        text_x = int(white_rect_width//2 - (font.getlength(TAG)) // 2)
        text_y = int(qr_y + qr_size  + 0 * ((white_rect_height - qr_y - qr_size - new_height/scale_factor) // 3))

        image.paste(resized_text, (text_x, text_y), resized_text)



        # Position text below QR Code
        text_x = 0 #qr_x + (qr_size - font.getlength(TAG)) // 2
        text_y = qr_y + qr_size - 40 + 0 * ((white_rect_height - qr_y - qr_size - text_height) // 3)
        #draw.text((text_x, text_y), TAG, fill=TEXT_COLOR, font=font) Removido por uso de Step 3



        '''#pasar la imagen a un tamaño adecuado para la impresión de la etiqueta en matriz insertada en tabla en Microsoft Word
        medida_image_MSWord_inch = 4.30 / 2.54
        dpi_value = BASE_WIDTH/medida_image_MSWord_inch
        #image.save(output_path, dpi=(dpi_value, dpi_value))
        print(f"Image saved as {output_path}")'''


    except IOError:
        print("Error: Arial font not found. Using default font.")
        font = ImageFont.load_default()
        text_x = qr_x
        text_y = qr_y + qr_size + 10
        draw.text((text_x, text_y), TAG, fill=TEXT_COLOR, font=font)
    return image

def create_qr_with_logo_label_and_frame(url, logo_path, output_path, qr_size, label="BORRAR", logo_size_ratio=0.2, frame_thickness=10,font_path="arialbd.ttf", font_size=48):
    # Create a QR Code instance
    qr = qrcode.QRCode(
        version=1,  # Controls the size of the QR Code. Higher number = larger code.
        error_correction=qrcode.constants.ERROR_CORRECT_H,  # High error correction for logo.
        box_size=10,  # Size of each box in pixels.
        border=4,  # Width of the border (minimum is 4 for QR codes).
    )
    qr.add_data(url)
    qr.make(fit=True)

    # Create the Canvas image
    qr_img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
    qr_img =qr_img.resize((qr_size, qr_size))


    # Open the logo image
    logo = Image.open(logo_path)

    # Calculate the logo size
    logo_size = int(qr_size * logo_size_ratio)
    logo = logo.resize((logo_size, logo_size), Image.Resampling.LANCZOS)

    # Calculate position for the logo RPCI
    logo_position = (
        (qr_img.width - logo.width) // 2,
        (qr_img.height - logo.height) // 2
    )

    # Paste the logo onto the QR code
    qr_img.paste(logo, logo_position, mask=logo if "A" in logo.mode else None)
    label_height = frame_thickness + font_size + 10  # Extra space for label and padding
    return qr_img, qr_img.height, qr_img.width


logo_path = "Aros _RPCI.jpg"  # Path to your logo image file
font_path = "arialbd.ttf"  # Path to Arial Black font file on your system
#Para Equipos e Instrumentos
#BASE_WIDTH = 400 #valor pendiente de revisión
#Para Blower y CCM
BASE_WIDTH = 800 #valor pendiente de revisión
BACKGROUND_COLOR = (255, 255, 255)  # White
TEXT_COLOR = (0, 0, 0)  # Black
# Calculate dimensions para tags tipo Tarjeta de Credito
white_rect_width = int(BASE_WIDTH/1.64) # 1.64 valor obtenido de dividir 4.30cm/2.54 inchs y 1.69 de la proporcion Ancho/Alto de la Credit Card
#white_rect_height = int(1.10 * BASE_WIDTH)  #Height es 10% mayor que el Width
white_rect_height = int(BASE_WIDTH/1.64/1.69)  # 1.64 valor obtenido de dividir 4.30cm/2.54 inchs y 1.69 de la proporcion Ancho/Alto de la Credit Card
top_margin = int(0.01 * white_rect_height)
bottom_margin = int(0.05 * white_rect_height)
qr_size = int(0.80 * white_rect_height)  # QR replaces pink square

#Bloque With  para creación de Texto del TAG debajo del QR
with open("TAGS.csv", mode='r', newline='', encoding='utf-8') as file:
    reader = csv.DictReader(file,delimiter=";")  # Use DictReader to access columns by name

    # Iterate through each row
    for row in reader:
        TAG = row['TAG']  # Access the TAG column
        PREFIX=row['PREFIX']
        LINK = row['LINK']  # Access the Link column
        output_path = f'URLS/{TAG}.png'  # Output file path for the QR code with logo, label, and frame

        #Function for create QR Image for each TAG with middle logo in QR image
        qr_image, qr_height, qr_width = create_qr_with_logo_label_and_frame(LINK, logo_path, output_path,qr_size, label=TAG, font_path=font_path)

        qr_image_with_TAG =create_TagTex_at_Bottom(qr_image)


        final_image_cardsize = create_tag_text_logoRPCI(qr_image_with_TAG,int(qr_height),int(qr_width))

        # Save image
        output_path = output_path = f'URLS/{TAG}.png'
        final_image_cardsize.save(output_path)
#Lllevar Imagenes a Microsoft Word
final = create_WordDocument()
