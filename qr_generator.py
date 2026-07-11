"""Motor de generación de etiquetas QR (TCard y Adhesivo).

Este módulo consolida la lógica que antes estaba duplicada en `main.py`
(TCard) y `main_blowers.py` (Adhesivo). A diferencia de esos scripts, aquí
NO se ejecuta nada al importar: todo está dentro de funciones y se dispara
llamando a `generate(...)`. Esto permite usar el motor desde la GUI (gui.py),
desde otros scripts o desde la línea de comandos.

Las únicas diferencias reales entre los dos tipos de etiqueta son un puñado
de constantes; se recogen en PRESETS más abajo.
"""

import os
import csv

import qrcode
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# --------------------------------------------------------------------------- #
# Presets: lo único que diferencia una TCard de un Adhesivo.
# --------------------------------------------------------------------------- #
PRESETS = {
    "tcard": {
        "label": "TCard (Equipos e Instrumentos)",
        "orientation": "portrait",       # Word en vertical
        "word_img_w": 3.3,               # ancho de la imagen dentro de la tabla Word (in)
        "word_img_h": 2.16,              # alto de la imagen dentro de la tabla Word (in)
        "spacer_pt": 3,                  # tamaño de fuente de la fila separadora (pt)
        "rpci_logo_size": (168, 200),    # tamaño del logo RPCI (px)
        "logo_x_offset": 10,             # desplazamiento X del logo RPCI (px)
        "init_font_size": 0.5,           # tamaño de fuente inicial del texto del TAG
        "set_margins": True,             # fijar márgenes de página
        "set_col_widths": True,          # fijar ancho de columnas de la tabla
    },
    "adhesive": {
        "label": "Adhesivo (Blowers / CCM)",
        "orientation": "landscape",      # Word en horizontal
        "word_img_w": 4.33,              # 11 cm de ancho (tamaño físico real de la etiqueta)
        "word_img_h": None,              # alto automático: python-docx preserva la proporción del QR (~6.5 cm)
        "spacer_pt": 7,
        "rpci_logo_size": (181, 215),
        "logo_x_offset": 15,
        "init_font_size": 1,
        "set_margins": False,
        "set_col_widths": False,
    },
}

# Valores compartidos por ambos modos.
BASE_WIDTH = 800  # Súbelo si las imágenes salen pixeladas en el PDF.
BACKGROUND_COLOR = (255, 255, 255)
TEXT_COLOR = (0, 0, 0)

# Nombres de columna esperados en el CSV (delimitador ';').
REQUIRED_COLUMNS = ("TAG", "LINK")

# Ancho máximo (cm) de un adhesivo para caber dos por fila en Letter horizontal.
# Si el ancho pedido supera este valor, la tabla usa una sola columna (incluso
# con la hoja en horizontal).
ADHESIVE_TWO_COLUMN_MAX_CM = 13.0


def resolve_font_path():
    """Devuelve la primera fuente en negrita disponible según el sistema.

    USUARIOS DE WINDOWS: en Windows se usa Arial Bold ('C:/Windows/Fonts/
    arialbd.ttf'), que existe de forma nativa. En Linux/Mac esa ruta NO existe,
    así que se cae a DejaVu Sans Bold (equivalente). Si no se encuentra ninguna,
    se devuelve None y PIL usará su fuente por defecto (mucho más pequeña).
    """
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf",                              # Windows – Arial Bold
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",      # Debian / Ubuntu
        "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",              # Fedora / RHEL
        "/Library/Fonts/Arial Bold.ttf",                            # macOS
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",       # macOS
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def label_aspect_ratio():
    """Proporción ancho/alto de la etiqueta generada (fija por diseño).
    La GUI la usa para calcular el alto proporcional a partir del ancho."""
    d = _compute_dimensions()
    return d["white_rect_width"] / d["white_rect_height"]


def _compute_dimensions():
    """Dimensiones del lienzo tipo tarjeta de crédito, derivadas de BASE_WIDTH."""
    white_rect_width = int(BASE_WIDTH / 1.64)
    white_rect_height = int(BASE_WIDTH / 1.64 / 1.69)
    top_margin = int(0.01 * white_rect_height)
    qr_size = int(0.80 * white_rect_height)
    return {
        "white_rect_width": white_rect_width,
        "white_rect_height": white_rect_height,
        "top_margin": top_margin,
        "qr_size": qr_size,
    }


def create_qr_with_logo(link, client_logo_path, qr_size, logo_size_ratio=0.2):
    """Crea el QR a partir de `link` y le pega el logo del cliente en el centro."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,  # alta corrección para el logo
        box_size=10,
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    qr_img = qr_img.resize((qr_size, qr_size))

    logo = Image.open(client_logo_path)
    logo_size = int(qr_size * logo_size_ratio)
    logo = logo.resize((logo_size, logo_size), Image.Resampling.LANCZOS)

    logo_position = (
        (qr_img.width - logo.width) // 2,
        (qr_img.height - logo.height) // 2,
    )
    qr_img.paste(logo, logo_position, mask=logo if "A" in logo.mode else None)
    return qr_img


def create_tag_text_at_bottom(qr_image, tag, cfg, dims, font_path):
    """Pega el QR en el lienzo y renderiza el texto del TAG debajo, ajustando
    el tamaño de fuente hasta que el ancho del texto encaje con el del QR."""
    white_rect_width = dims["white_rect_width"]
    white_rect_height = dims["white_rect_height"]
    top_margin = dims["top_margin"]
    qr_size = dims["qr_size"]

    image = Image.new("RGB", (white_rect_width, white_rect_height), BACKGROUND_COLOR)
    draw = ImageDraw.Draw(image)

    qr_x = 0
    qr_y = top_margin
    image.paste(qr_image, (qr_x, qr_y))

    if font_path is None:
        # Sin fuente TrueType disponible: usar la fuente por defecto de PIL.
        font = ImageFont.load_default()
        draw.text((qr_x, qr_y + qr_size + 10), tag, fill=TEXT_COLOR, font=font)
        return image

    font_size = cfg["init_font_size"]

    # Busca el tamaño de fuente óptimo para que el ancho del texto encaje.
    while True:
        font = ImageFont.truetype(font_path, font_size)
        text_width = font.getlength(tag)
        if text_width > BASE_WIDTH / 1.64 * 1.1 and len(tag) > 12:
            break
        elif text_width > BASE_WIDTH / 1.64 * 0.85 and len(tag) > 6:
            break
        elif text_width > BASE_WIDTH / 1.64 * 0.4 and len(tag) <= 6:
            break
        font_size += 1

    font_size = max(font_size - 1, 1)
    font = ImageFont.truetype(font_path, font_size)

    # Paso 1: alto inicial del texto.
    text_bbox = font.getbbox(tag)
    text_width = text_bbox[2] - text_bbox[0]
    if len(tag) > 12:
        text_height = int((text_bbox[3] - text_bbox[1]) * 2)
    else:
        text_height = int((text_bbox[3] - text_bbox[1]) * 1)

    text_image = Image.new("RGBA", (text_width, text_height), (0, 0, 0, 0))
    text_draw = ImageDraw.Draw(text_image)
    text_draw.text((-text_bbox[0], -text_bbox[1]), tag, font=font, fill="black")

    # Paso 2: reducir el alto (manteniendo el ancho) para textos cortos.
    scale_factor = 0.65 if len(tag) < 12 else 1
    new_height = int(text_height * scale_factor)
    resized_text = text_image.resize(
        (text_width, new_height),
        resample=Image.Resampling.LANCZOS,
    )

    # Paso 3: pegar el texto centrado bajo el QR.
    text_x = int(white_rect_width // 2 - font.getlength(tag) // 2)
    text_y = int(qr_y + qr_size)
    image.paste(resized_text, (text_x, text_y), resized_text)

    return image


def create_tag_text_logo_rpci(image, cfg, dims, rpci_logo_path):
    """Pega el logo RPCI a la derecha del QR / encima del texto del TAG."""
    white_rect_width = dims["white_rect_width"]
    top_margin = dims["top_margin"]

    logo = Image.open(rpci_logo_path)
    resized_logo = logo.resize(cfg["rpci_logo_size"])

    logo_x = white_rect_width / 2 + cfg["logo_x_offset"]
    logo_y = top_margin + 5

    image.paste(resized_logo, (int(logo_x), int(logo_y)))
    return image


def _size_by_width(image_path, width_cm):
    """Devuelve (width_in, height_in) para que la imagen tenga EXACTAMENTE el
    ancho pedido (en cm) y el alto se calcule por proporción, sin deformar el
    QR ni el logo. El ancho manda; el alto es consecuencia."""
    with Image.open(image_path) as im:
        px_w, px_h = im.size
    img_aspect = px_w / px_h
    width_in = width_cm / 2.54
    return width_in, width_in / img_aspect


def create_word_document(image_folder, docx_path, cfg, label_width_cm=None, columns=2):
    """Lee las imágenes de `image_folder` y las coloca en una tabla dentro de
    un documento Word, según el preset activo.

    Si `label_width_cm` es un número, cada imagen se inserta con EXACTAMENTE ese
    ancho en cm y el alto se calcula por proporción (sin deformar el QR/logo).
    Si es None, se usa el tamaño fijo del preset (`word_img_w`/`word_img_h`).

    `columns` es el número de imágenes por fila (por defecto 2). Para adhesivos
    anchos (>13 cm) se usa 1 columna, ya que dos no caben en el ancho de la
    hoja Letter en horizontal."""
    columns = max(1, int(columns))
    image_files = [
        os.path.join(image_folder, f)
        for f in sorted(os.listdir(image_folder))
        if f.lower().endswith(("png", "jpg", "jpeg"))
    ]

    doc = Document()
    section = doc.sections[0]

    if cfg["set_margins"]:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    if cfg["orientation"] == "landscape":
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    table = doc.add_table(rows=0, cols=columns)

    if cfg["set_col_widths"]:
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(5)

    img_w = Inches(cfg["word_img_w"])
    # Si el alto es None, se pasa solo el ancho y python-docx calcula el alto
    # preservando la proporción original de la imagen (sin deformar el QR).
    img_h = Inches(cfg["word_img_h"]) if cfg["word_img_h"] is not None else None

    def _add_picture(paragraph, path):
        run = paragraph.add_run()
        if label_width_cm is not None:
            # El usuario definió el ancho en cm: se respeta el ancho y el alto
            # se calcula por proporción (sin deformar).
            w_in, h_in = _size_by_width(path, label_width_cm)
            run.add_picture(path, width=Inches(w_in), height=Inches(h_in))
        elif img_h is None:
            run.add_picture(path, width=img_w)
        else:
            run.add_picture(path, width=img_w, height=img_h)

    for i in range(0, len(image_files), columns):
        row = table.add_row().cells

        for col in range(columns):
            idx = i + col
            if idx >= len(image_files):
                break
            paragraph = row[col].paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _add_picture(paragraph, image_files[idx])

        # Fila separadora en blanco (controla el espaciado vertical).
        blank_row = table.add_row().cells
        for cell in blank_row:
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(" ")
            run.font.size = Pt(cfg["spacer_pt"])

    doc.save(docx_path)
    return docx_path


def read_tags(csv_path):
    """Lee el CSV (delimitador ';') y devuelve la lista de filas.

    Lanza ValueError si faltan las columnas TAG o LINK.
    """
    with open(csv_path, mode="r", newline="", encoding="utf-8-sig") as file:
        reader = csv.DictReader(file, delimiter=";")
        fieldnames = reader.fieldnames or []
        missing = [c for c in REQUIRED_COLUMNS if c not in fieldnames]
        if missing:
            raise ValueError(
                f"El CSV no tiene las columnas requeridas: {', '.join(missing)}. "
                f"Columnas encontradas: {', '.join(fieldnames)}"
            )
        return [row for row in reader if (row.get("TAG") or "").strip()]


def generate(
    csv_path,
    mode="tcard",
    output_dir="URLS",
    docx_path="Images_Table.docx",
    client_logo_path="cliente.png",
    rpci_logo_path="LOGO_RPCI.jpg",
    width_cm=None,
    height_cm=None,
    progress=None,
):
    """Genera todas las etiquetas QR y el documento Word.

    Parámetros
    ----------
    csv_path : ruta al CSV fuente (delimitador ';', columnas DOMAIN;SUBSITE;TAG;LINK).
    mode : 'tcard' o 'adhesive'.
    output_dir : carpeta donde se guardan los PNG (se crea si no existe).
    docx_path : ruta del documento Word de salida.
    client_logo_path : logo del cliente que va en el centro del QR.
    rpci_logo_path : logo RPCI que va a la derecha del QR.
    width_cm : si se indica, cada imagen del Word se inserta con EXACTAMENTE ese
        ancho en cm y el alto se calcula por proporción (sin deformar el QR/logo).
        El ancho manda. Si no se indica, se usa el tamaño fijo del preset.
    height_cm : ignorado para el tamaño (el alto se deriva del ancho por
        proporción); se conserva por compatibilidad con la GUI.
    progress : callback opcional progress(done, total, message) para la GUI.

    Devuelve un dict con {'tags': [...], 'output_dir', 'docx_path'}.
    """
    if mode not in PRESETS:
        raise ValueError(f"Modo desconocido '{mode}'. Usa uno de: {', '.join(PRESETS)}")
    cfg = PRESETS[mode]

    for path, name in [(csv_path, "CSV"), (client_logo_path, "logo del cliente"),
                       (rpci_logo_path, "logo RPCI")]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"No se encontró el archivo de {name}: {path}")

    os.makedirs(output_dir, exist_ok=True)  # corrige el bug: antes URLS/ debía existir

    dims = _compute_dimensions()
    font_path = resolve_font_path()

    rows = read_tags(csv_path)
    total = len(rows)
    if total == 0:
        raise ValueError("El CSV no contiene filas con TAG.")

    def _report(done, message):
        if progress:
            progress(done, total, message)

    generated_tags = []
    for idx, row in enumerate(rows, start=1):
        tag = row["TAG"].strip()
        link = (row.get("LINK") or "").strip()
        _report(idx - 1, f"Generando {tag} ({idx}/{total})")

        qr_image = create_qr_with_logo(link, client_logo_path, dims["qr_size"])
        image = create_tag_text_at_bottom(qr_image, tag, cfg, dims, font_path)
        image = create_tag_text_logo_rpci(image, cfg, dims, rpci_logo_path)

        image.save(os.path.join(output_dir, f"{tag}.png"))
        generated_tags.append(tag)
        _report(idx, f"Generado {tag} ({idx}/{total})")

    label_width_cm = None
    columns = 2
    if width_cm:
        label_width_cm = float(width_cm)
        # Un adhesivo más ancho que el umbral no cabe dos veces: se usa una
        # sola columna incluso con la hoja en horizontal.
        if label_width_cm > ADHESIVE_TWO_COLUMN_MAX_CM:
            columns = 1

    _report(total, "Creando documento Word…")
    create_word_document(output_dir, docx_path, cfg,
                         label_width_cm=label_width_cm, columns=columns)
    _report(total, "¡Listo!")

    return {"tags": generated_tags, "output_dir": output_dir, "docx_path": docx_path}


if __name__ == "__main__":
    # Uso por línea de comandos:
    #   python qr_generator.py tcard TAGS.csv
    #   python qr_generator.py adhesive TAGS.csv
    import sys

    mode_arg = sys.argv[1] if len(sys.argv) > 1 else "tcard"
    csv_arg = sys.argv[2] if len(sys.argv) > 2 else "TAGS.csv"

    def _cli_progress(done, total, message):
        print(f"[{done}/{total}] {message}")

    result = generate(csv_arg, mode=mode_arg, progress=_cli_progress)
    print(f"Generadas {len(result['tags'])} etiquetas → {result['docx_path']}")
